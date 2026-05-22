"""
generate_test_formats.py
从 TXT 源生成 PDF 和 DOCX 测试文件（无需手动操作 Word）
"""
from pathlib import Path

SOURCE_TXT = "data/source_docs/deposit_insurance_regulation.txt"
OUT_DIR = "data/source_docs"

# ---------- 生成 DOCX ----------
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()
    # 标题
    title = doc.add_heading("存款保险条例", level=0)
    # 副标题
    doc.add_paragraph(
        "2014年10月29日国务院第67次常务会议通过\n"
        "2015年2月17日中华人民共和国国务院令第660号公布\n"
        "自2015年5月1日起施行"
    )

    # 读取 TXT 并逐条写入
    text = Path(SOURCE_TXT).read_text(encoding="utf-8")
    # 跳过前 5 行（标题和副标题已手动写入）
    body = text.split("\n", 5)[-1] if len(text.split("\n")) > 5 else text

    for line in body.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("第") and "条" in line[:5]:
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    docx_path = f"{OUT_DIR}/deposit_insurance_regulation.docx"
    doc.save(docx_path)
    print(f"✅ DOCX 已生成: {docx_path}")

except ImportError:
    print("⚠️ 需要 python-docx: pip install python-docx")

# ---------- 生成 PDF ----------
try:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # 查找系统可用中文字体
    font_candidates = [
        "C:/Windows/Fonts/simhei.ttf",   # 黑体
        "C:/Windows/Fonts/msyh.ttc",     # 微软雅黑
        "C:/Windows/Fonts/simsun.ttc",   # 宋体
    ]
    font_path = None
    for fp in font_candidates:
        if Path(fp).exists():
            font_path = fp
            break

    if font_path is None:
        print("❌ 未找到任何中文字体文件，PDF 将无法显示中文")
    else:
        pdf.add_font("Chinese", fname=font_path)
        pdf.set_fallback_fonts(["Chinese"])   # 备用字体，防止漏字

        pdf.set_font("Chinese", size=16)
        pdf.cell(0, 10, "存款保险条例", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
        pdf.set_font("Chinese", size=10)
        pdf.cell(0, 6, "国务院令第660号", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
        pdf.ln(5)

        for line in body.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("第") and "条" in line[:5]:
                pdf.set_font("Chinese", size=11)
                pdf.ln(2)
                pdf.cell(0, 6, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.set_font("Chinese", size=10)
            else:
                pdf.multi_cell(0, 5, line)

        pdf_path = f"{OUT_DIR}/deposit_insurance_regulation.pdf"
        pdf.output(pdf_path)
        print(f"✅ PDF 已生成: {pdf_path}")

except ImportError:
    print("⚠️ 需要 fpdf2: pip install fpdf2")